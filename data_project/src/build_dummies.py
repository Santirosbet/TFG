# -*- coding: utf-8 -*-
"""
Generates output/TFG for Dummies.docx
A plain-English study guide for thesis defense preparation.
Run: python src/build_dummies.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "output", "TFG for Dummies.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1f, 0x4e, 0x79)
ORANGE = RGBColor(0xff, 0x6b, 0x35)
GREEN  = RGBColor(0x38, 0x86, 0x38)
GREY   = RGBColor(0x55, 0x55, 0x55)
BLACK  = RGBColor(0x00, 0x00, 0x00)
RED    = RGBColor(0xc0, 0x00, 0x00)


# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ORANGE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
    return p

def analogy(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run("ANALOGY: ")
    r1.bold = True
    r1.font.color.rgb = GREEN
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREEN
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run("DEFENSE TIP: ")
    r1.bold = True
    r1.font.color.rgb = RED
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RED
    return p

def keynum(label, value, explain):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(11)
    r2 = p.add_run(f"{value}  ")
    r2.bold = True
    r2.font.color.rgb = ORANGE
    r2.font.size = Pt(11)
    r3 = p.add_run(f"— {explain}")
    r3.font.size = Pt(11)
    r3.font.color.rgb = GREY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def qa(question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"Q: {question}")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.0)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(f"A: {answer}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLACK

def pagebreak():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("─" * 70)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("TFG FOR DUMMIES")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Predictive Analysis of Respiratory Medicine Inventory Demand")
r2.font.size = Pt(16)
r2.font.color.rgb = ORANGE
r2.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(10)
r3 = p3.add_run("Your plain-English guide to understanding, explaining\nand defending every part of your thesis")
r3.font.size = Pt(13)
r3.font.color.rgb = GREY

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(30)
r4 = p4.add_run("Santiago Rosales Betancourt  |  TFG — Business Intelligence & Data Analytics")
r4.font.size = Pt(11)
r4.font.color.rgb = GREY

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# THE BIG PICTURE (start here)
# ══════════════════════════════════════════════════════════════════════════════
h1("THE BIG PICTURE — Start Here")
body(
    "Imagine you run the supply chain for a large pharmacy chain. Every year, you need "
    "to order extra respiratory medicines (inhalers, antihistamines) before winter. "
    "The problem: manufacturing lead times are 6-8 months. If you order in January "
    "when you already know it's a bad flu season, it's too late — the factory needs "
    "the order by July."
)
body(
    "So the real question is: can you know in July whether the coming January will be "
    "a bad flu season or a mild one? That is exactly what this thesis answers."
)
analogy(
    "It's like a weather forecast that works 6 months in advance. Not perfect, "
    "but much better than guessing — and enough to save real money."
)
body(
    "The answer: YES, and the key is Australia. The flu season in Australia peaks in "
    "June-July (their winter). That information is available in real-time via WHO data. "
    "And statistically, a bad Australian flu season in July predicts a bad European "
    "flu season in January — 28 weeks later. That 28-week window is longer than the "
    "typical manufacturing lead time for medicines."
)
h2("The One-Sentence Summary")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r = p.add_run(
    '"Using free, real-time WHO flu data from Australia, this thesis builds a machine '
    'learning model that predicts European respiratory medicine demand 26 weeks in '
    'advance, enabling pharmaceutical supply chains to order stock before it\'s needed."'
)
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# KEY NUMBERS CHEAT SHEET
# ══════════════════════════════════════════════════════════════════════════════
h1("KEY NUMBERS — Memorise These")
body("These are the numbers that will come up in every question. Know them cold.")

keynum("r = 0.70", "p < 0.001", "Cross-correlation between Australian flu and European R03 demand at lag 28 weeks. Very strong and statistically significant.")
keynum("Lag", "26-28 weeks", "How far ahead Australia predicts Europe. ~6-7 months. Longer than manufacturing lead time.")
keynum("Model A MAPE", "46.45%", "Autoregressive only (no flu data). The baseline you beat.")
keynum("Model B MAPE", "44.16%", "XGBoost + Australian flu signal. Test set (60 weeks).")
keynum("Switching Rule MAPE", "35.78%", "Best result in the project. Hybrid: XGBoost in winter, historical mean in summer.")
keynum("DM test", "6.23, p < 0.001", "Diebold-Mariano: XGBoost is statistically significantly better than SARIMA.")
keynum("Walk-forward", "192 predictions, 48 folds", "Robust validation. Model B wins in 69% of folds.")
keynum("Inventory saving", "EUR 273", "XGBoost saves this over 60 weeks vs naive policy. One pharmacy. Scales to millions.")
keynum("Countries validated", "7", "Australia, NZ, Chile, Argentina, Brazil, South Africa, Uruguay.")
keynum("R06 MAPE", "52.07%", "Framework works for antihistamines too, not just R03.")
keynum("CI coverage", "74% (80% nominal)", "Confidence intervals are well-calibrated (within 6pp of nominal).")

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 1: Introduction — Why Does This Thesis Exist?")

h2("1.1 Context and Motivation")
body(
    "The pharmaceutical supply chain is broken in a specific way: it reacts to demand "
    "instead of anticipating it. This creates the Bullwhip Effect — small demand "
    "fluctuations at the pharmacy level become massive order swings at the factory level. "
    "The result: either stockouts (people can't get medicines) or excess inventory "
    "(money locked up in unsold stock)."
)
analogy(
    "Think of a game of Chinese Whispers. By the time the message (demand signal) "
    "reaches the factory, it has been amplified and distorted. The Bullwhip Effect "
    "is that amplification in a supply chain."
)
body(
    "The seasonal nature of respiratory medicines makes this worse. R03 (bronchodilators, "
    "inhalers) and R06 (antihistamines) both spike every January-February during the "
    "European flu season. Manufacturers know it will happen — they just don't know how bad "
    "it will be until it's already too late to order more."
)

h2("1.2 Problem Statement")
body(
    "The core problem: there is no reliable early-warning system for European pharmaceutical "
    "demand spikes. Existing methods are either too simple (naive averages) or too reactive "
    "(SARIMA models that only look at past demand)."
)
body(
    "The thesis proposes that epidemiological data from the Southern Hemisphere — "
    "specifically, WHO FluNet surveillance counts from Australia — can act as a leading "
    "indicator that bridges this 6-month forecasting gap."
)
tip(
    'If asked "why this topic?": say the gap between manufacturing lead times (6-8 months) '
    'and forecast horizon of traditional models (4-8 weeks) creates a structural blind spot '
    'in pharmaceutical supply chains. This thesis plugs that gap using publicly available data.'
)

h2("1.3 Objectives")
bullet("Quantify the statistical lead-lag relationship between Southern Hemisphere flu activity and European pharmaceutical demand")
bullet("Build a machine learning forecasting model that exploits this relationship")
bullet("Validate the model against statistical baselines using rigorous tests (DM test, walk-forward validation)")
bullet("Translate the model into an operational tool (inventory simulation + Streamlit dashboard)")
bullet("Generalise the framework across 7 SH countries and 2 drug categories (R03 + R06)")

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 2: Literature Review — What Did Other People Find?")
body(
    "This chapter shows you read the academic literature. It's not about what YOU did "
    "— it's about what OTHERS have done and why there's still a gap for your work to fill."
)

h2("2.1 Pharmaceutical Inventory Management")
body(
    "Existing literature shows most pharma inventory models use simple time-series "
    "methods (moving averages, SARIMA). They are good at capturing past patterns but "
    "cannot predict structural demand changes caused by epidemic severity. "
    "Key references show the industry knows the problem but has no data-driven solution."
)
analogy(
    "These models are like driving a car by only looking in the rearview mirror. "
    "Fine for straight roads, useless when there's a sharp bend ahead."
)

h2("2.2 Hemispheric Epidemiology and the Lead-Lag")
body(
    "Several epidemiological studies confirm that flu waves travel globally in a "
    "predictable seasonal pattern. The Southern Hemisphere winter (June-August) "
    "consistently precedes the Northern Hemisphere winter (December-February) by "
    "approximately 6 months. This is the epidemiological basis for the entire thesis."
)
body(
    "Prior work by Viboud et al. (2006) and others documented the global spread pattern. "
    "However, no prior study had specifically quantified this relationship as a "
    "pharmaceutical demand forecasting signal — that is the novel contribution."
)

h2("2.3 Machine Learning for Demand Forecasting")
body(
    "XGBoost (Chen & Guestrin, 2016) has become the standard benchmark for tabular "
    "prediction tasks. Studies in retail and healthcare show it consistently outperforms "
    "SARIMA and ARIMA for seasonal demand with external covariates. SHAP values "
    "(Lundberg & Lee, 2017) provide interpretability, which is critical for "
    "regulatory acceptance in healthcare applications."
)

h2("2.4 Research Gap")
body(
    "The gap: no study had combined (1) real WHO FluNet hemispheric flu surveillance data, "
    "(2) real pharmaceutical sales data, and (3) an XGBoost model with epidemiological "
    "lag features, validated with walk-forward CV and compared against SARIMA using a "
    "formal Diebold-Mariano test. This thesis fills all four simultaneously."
)
tip(
    'If asked "what is your contribution vs existing literature?": '
    '"Previous studies documented the epidemiological lead-lag pattern. My contribution '
    'is translating that pattern into a validated, operational forecasting pipeline with '
    'formal statistical tests and a real-time decision dashboard."'
)

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: STATE OF THE ART
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 3: State of the Art — What Tools Exist?")
body(
    "This chapter reviews the current state of the technical tools and frameworks used "
    "in the thesis. Think of it as the toolbox review before building something."
)

h2("3.1 Pharmaceutical Supply Chain Challenges")
body(
    "Key challenges: demand volatility, long lead times, regulatory constraints on "
    "batch production, cold chain requirements, and the Bullwhip Effect. The literature "
    "consistently shows that reactive ordering amplifies volatility through the chain."
)

h2("3.2 Seasonality and the Lead-Lag Effect")
body(
    "Seasonal time series are well-studied. The SARIMA family handles seasonality via "
    "seasonal differencing. The novel claim in this thesis is that an EXTERNAL seasonal "
    "signal (AU flu) can predict European demand better than the endogenous seasonal "
    "pattern alone. This is the exogenous regressor hypothesis."
)

h2("3.3 Time Series Models Compared")
body("The three models compared in this thesis, explained simply:")
bullet("SARIMA: 'Smart moving average' that learns seasonal patterns from past data only. No external information. Like predicting weather from your own thermometer history.")
bullet("XGBoost Model A: Decision trees trained on past demand + time features. Faster and more flexible than SARIMA, but still endogenous only.")
bullet("XGBoost Model B: Same as A but adds the Australian flu signal as a feature. This is where the forward-looking capability comes from.")

h2("3.4 The Bullwhip Effect")
body(
    "When a pharmacy orders a little extra because they're worried about stock, the "
    "distributor gets nervous and orders even more from the wholesaler, who orders "
    "even more from the manufacturer. By the time the signal reaches the factory, "
    "a 10% demand increase at the pharmacy level looks like a 40% increase. "
    "This amplification is the Bullwhip Effect."
)
analogy(
    "Like a game of telephone, but everyone gets progressively more panicked. "
    "The original message was 'sell a few more units'. The factory hears 'emergency, "
    "double production'. Accurate forecasting breaks this chain of panic."
)

h2("3.5 XGBoost vs SARIMA: Why XGBoost?")
body(
    "SARIMA is the classical statistical tool. XGBoost is the modern ML tool. "
    "XGBoost wins because: (1) it handles non-linear relationships between features, "
    "(2) it can incorporate external variables (flu signals) naturally, "
    "(3) it is more flexible with irregular patterns. The thesis formally proves "
    "XGBoost wins with DM statistic = 6.23, p < 0.001."
)

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 4: Methodology — How Did You Do It?")
body(
    "This chapter explains the research design and every step of the data pipeline. "
    "Think of it as the recipe: ingredients, tools, and instructions."
)

h2("4.1 Data Sources and Validation")
body("Four datasets, each playing a specific role:")
bullet("WHO FluNet: 183,732 weekly flu observations, 140+ countries, 1997-2024. FREE, public API. This is the key epidemiological signal.")
bullet("Kaggle EU Pharma Sales: 302 weeks (2014-2019) from one European pharmacy. This is your TARGET VARIABLE — what you're trying to predict.")
bullet("AMELI France Open Medic: Annual national dispensation data for France. Used to VALIDATE that the Kaggle dataset is representative of European patterns.")
bullet("PBS Australia: 50,979 monthly prescription records 2020-2024. Used as additional Australian signal for the post-2019 period.")

analogy(
    "FluNet is the thermometer (input), Kaggle is the patient's temperature chart "
    "(output you want to predict), AMELI is a second opinion from a different doctor "
    "(validation), PBS is a more detailed Australian reading."
)

h2("4.2 Data Cleaning")
body(
    "All four datasets had different formats, frequencies, and date conventions. "
    "Key cleaning steps: (1) resample all data to weekly frequency aligned to ISO "
    "week calendar; (2) handle missing values in FluNet sparse country data; "
    "(3) aggregate Australian state data to national level; (4) create lagged features "
    "(shift Australian signal 26 weeks forward to align with EU peak)."
)

h2("4.3 Critical Methodological Decision: No Synthetic Proxy")
body(
    "An early version of the model used a synthetic proxy for Australian flu based on "
    "European data (circular reasoning: using EU data to predict EU data). "
    "This was identified as a fundamental flaw and corrected: the model uses ONLY "
    "real Australian WHO FluNet data as the leading indicator. This correction is "
    "documented explicitly in Section 4.3 of the thesis."
)
tip(
    '"Why did you correct the proxy?" — Because using European data to predict European '
    'data is circular. It would inflate model performance artificially. Scientific '
    'integrity requires using genuinely independent data as the predictor.'
)

h2("4.4 Feature Engineering")
body("The final feature set for Model B (4 features):")
bullet("R03_lag1: Last week's actual demand. Most important predictor.")
bullet("R03_lag4_avg: Rolling 4-week average demand. Smoothed recent trend.")
bullet("flu_au_lagged: Australian FluNet positives shifted 26 weeks forward. The leading indicator.")
bullet("flu_eu_positives: Current European FluNet positives. Contemporaneous signal.")
bullet("week_of_year: ISO week number. Captures seasonality directly.")

h2("4.5 Train/Test Split — Preventing Data Leakage")
body(
    "The data was split chronologically: training on 2014-2018 (approx. 242 weeks), "
    "test on the final 60 weeks (approx. 2018-2019). "
    "CRITICAL: no random shuffling. For time series, shuffling is wrong — it would "
    "let the model 'see the future' during training. This is data leakage."
)
analogy(
    "Leakage is like studying with the exam answers in front of you. "
    "Your score looks great but you learned nothing. "
    "Chronological splitting means the model has never seen test week data during training."
)

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5: EDA
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 5: Exploratory Data Analysis — What Did the Data Show?")
body(
    "Before building any model, you explore the data visually and statistically "
    "to understand its patterns. This chapter answers: is the hypothesis plausible? "
    "Does the data actually show what we expect?"
)

h2("5.1 Cross-Correlation Analysis (CCF)")
body(
    "The CCF measures how correlated two time series are when one is shifted by N weeks. "
    "We computed the CCF between Australian FluNet positives and European R03 demand "
    "for all lags from -52 to +52 weeks."
)
body(
    "RESULT: The maximum correlation occurs at lag +28 weeks: r = 0.70, p < 0.001. "
    "This means Australian flu activity today has a 0.70 correlation with European "
    "medicine demand 28 weeks from now. This is the empirical foundation of the thesis."
)
analogy(
    "Imagine two runners in a relay race. Runner 1 (Australia) runs their lap first. "
    "The CCF measures how closely Runner 2's (Europe's) performance 28 weeks later "
    "mirrors Runner 1's speed. r=0.70 means they are very closely linked."
)
tip(
    '"Why 28 weeks specifically?" — The CCF peak. The mathematical maximum correlation '
    'falls at lag 28. We also validated at lag 26 (the optimal feature lag) because '
    'with real data, the 26-week lag gave better model performance even though the '
    'theoretical CCF peak is at 28. This small discrepancy is normal: population '
    'dynamics take time to translate into pharmacy-level purchasing.'
)

h2("5.2 Time-Series Splitting")
body(
    "The 302-week dataset was split: 80% train, 20% test. The split point was chosen "
    "to give exactly 60 test weeks — enough to cover a full seasonal cycle while "
    "maximising training data."
)

h2("5.3 Outlier Analysis")
body(
    "Some weeks had zero sales (pharmacy closed, data gaps) or extreme spikes. "
    "These were handled conservatively: genuine outliers (recording errors) were "
    "imputed with local means; genuine demand spikes (e.g., flu alerts) were kept "
    "because they represent real variation the model must learn to predict."
)

h2("5.4 Modelling Strategy")
body(
    "Two-phase approach: first establish a strong autoregressive baseline (Model A), "
    "then add external epidemiological features (Model B). All comparisons use identical "
    "train/test splits and hyperparameters to isolate the effect of the new features."
)

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6: MODEL DEVELOPMENT (the core)
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 6: Model Development — The Heart of the Thesis")
body(
    "This is the longest and most important chapter. It documents the full research "
    "journey from a simple baseline to the final optimised forecasting framework."
)

h2("6.0 Research Journey — The Story of the Model")
body(
    "The model did NOT emerge fully formed. It was built iteratively, adding one data "
    "source at a time, testing each hypothesis, and letting the results guide the "
    "next step. This is real research methodology."
)

h3("Phase 1: Kaggle Only (Starting Point)")
body("Model A — autoregressive baseline. Features: last week's demand + 4-week average + week_of_year.")
bullet("MAPE: 46.45%")
bullet("What it can do: predict the general seasonal shape")
bullet("What it CANNOT do: forecast peaks with any lead time — it only knows what happened last week")

h3("Phase 2: Adding Australian FluNet (The Key Insight)")
body("Model B — adds flu_au_lagged (Australian FluNet shifted 26 weeks).")
bullet("MAPE: 44.16% on 60-week test set")
bullet("DM test vs SARIMA: p < 0.001 (statistically significant win)")
bullet("Now has 26-week forward-looking capability — the core innovation")
analogy("Adding Australia's data to the model is like adding a weather radar to a ship's navigation system. Before, you were steering by memory of past storms. Now you can see the storm forming 6 months out.")

h3("Phase 3: Validating with Global FluNet + SHAP")
body("Added European FluNet as contemporaneous signal. SHAP analysis confirmed feature importance hierarchy.")
bullet("Feature importance: R03_lag1 (most) > flu_au_lagged > flu_eu_positives > week_of_year")
bullet("Walk-forward validation: 48 folds, 192 predictions, Model B wins 69% of folds")
bullet("WFV MAPE: 48.63% (higher than test MAPE because it includes more off-season folds)")

h3("Phase 4: French AMELI External Validation")
body("Used national French data to verify the Kaggle dataset is not an anomaly.")
bullet("Scale ratio: France ~1,000x one pharmacy (expected for 68M population)")
bullet("Seasonal pattern: identical peak in weeks 1-10 (Jan-Feb)")
bullet("Conclusion: Kaggle dataset is a valid representative sample of European R03 seasonality")

h3("Phase 5: 7-Country Southern Hemisphere Validation")
body("Tested the framework with AU, NZ, Chile, Argentina, Brazil, South Africa, Uruguay.")
bullet("Key finding: high CCF correlation does NOT guarantee low model MAPE")
bullet("NZ has lowest isolated MAPE (41.4%) despite weakest CCF (r=0.264)")
bullet("Australia remains recommended: best data quality, network size, consistency")

h3("Phase 6: Seasonal Switching Rule (Best Result)")
body("Identified that summer off-season (weeks 22-39) creates artificially high MAPE.")
bullet("Solution: use XGBoost in peak season, historical seasonal mean in off-season")
bullet("Result: MAPE drops from 48.63% to 35.78% — a 26.4% relative improvement")
bullet("This is the BEST result in the entire project")
tip(
    '"Why does the switching rule work?" — Because the model has fundamentally different '
    'error profiles in summer vs. winter. In summer, demand is near-zero and unpredictable '
    'with ML. The historical mean is provably better during those weeks. Recognising this '
    'and designing a hybrid solution is the methodological contribution of Section 6.13.'
)

h3("Phase 6b: Ensemble (Negative Result — Still Important)")
body("Tested AU+NZ, AU+Chile, AU+NZ+Chile combinations.")
bullet("All ensemble variants WORSE than Australia alone")
bullet("Australia alone MAPE: 52.52%; AU+NZ MAPE: 52.86%")
bullet("Conclusion: more signals ≠ better prediction at this data scale")
tip(
    '"Wasn\'t the ensemble a failure?" — No. A negative result is still a scientific '
    'contribution. It rules out a promising hypothesis and tells future researchers '
    'not to pursue this direction without more data. The thesis explicitly documents '
    'and explains this finding rather than hiding it.'
)

h2("6.1 XGBoost — How It Works (Simple Version)")
body(
    "XGBoost builds a series of decision trees, each one learning from the mistakes "
    "of the previous one. It's called 'boosting' — each tree 'boosts' the overall "
    "prediction by fixing remaining errors."
)
analogy(
    "Like a team of experts reviewing each other's work. Expert 1 makes a prediction. "
    "Expert 2 sees where Expert 1 was wrong and tries to fix it. Expert 3 fixes what "
    "Expert 2 got wrong. After 300 experts, the final combined prediction is very good."
)
body("Key hyperparameters (what they do):")
bullet("n_estimators = 300: number of trees (experts). More = more accurate but slower.")
bullet("max_depth = 4: how detailed each tree can be. Too deep = overfitting.")
bullet("learning_rate = 0.05: how fast each tree learns. Slow = more stable.")
bullet("subsample = 0.8: each tree sees only 80% of data. Prevents over-reliance on any one week.")

h2("6.7 SHAP Explainability")
body(
    "SHAP (SHapley Additive exPlanations) decomposes each prediction into contributions "
    "from each feature. If the model predicts 85 units this week, SHAP tells you: "
    "'R03_lag1 contributed +15, flu_au_lagged contributed +12, week_of_year contributed +8...'"
)
analogy(
    "Imagine 4 people sharing a taxi fare. SHAP tells you exactly how much each person "
    "should pay based on how far out of the way the taxi went for each of them. "
    "It's a fair attribution system for ML predictions."
)
tip(
    '"Why do you need SHAP if the model already works?" — Regulatory and managerial '
    'trust. A black-box model that says "trust me" cannot be deployed in a regulated '
    'pharmaceutical environment. SHAP provides the audit trail: you can show a manager '
    'WHY the model predicted a high-demand week, not just THAT it predicted one.'
)

h2("6.8 Walk-Forward Validation")
body(
    "Standard train/test split evaluates the model once. Walk-forward validation (WFV) "
    "simulates real deployment: train on weeks 1-200, predict weeks 201-204; "
    "train on weeks 1-204, predict weeks 205-208; and so on. 48 iterations."
)
body("This is the gold standard for time series validation. Results:")
bullet("192 total out-of-sample predictions")
bullet("Model B beats Model A in 69% of folds")
bullet("WFV MAPE: 48.63% (honest average across all seasons including summer)")
analogy(
    "Instead of one exam at the end of the year, you take 48 weekly quizzes — "
    "always on material you haven't studied yet. Much harder to cheat (overfit) "
    "and much more reliable evidence that the model actually works."
)

h2("6.9 SARIMA Baseline + Ljung-Box Test")
body(
    "SARIMA is the traditional statistical model for seasonal time series. "
    "It was fitted as an honest benchmark to compare against XGBoost."
)
body("Results:")
bullet("SARIMA MAPE: 56.81% (vs XGBoost 44.16%)")
bullet("XGBoost wins by 12.65 percentage points")
bullet("Diebold-Mariano statistic: 6.23, p < 0.001 (statistically significant)")
body(
    "The Ljung-Box test checked whether SARIMA residuals are white noise "
    "(properly modelled, no remaining patterns). Q(10) = 10.25, p = 0.419: "
    "PASSED. This means SARIMA is a well-specified baseline — XGBoost is not "
    "beating a broken model."
)
tip(
    '"Why test the SARIMA residuals?" — To prove you\'re comparing fairly. '
    'If SARIMA was badly specified, beating it would mean nothing. '
    'The Ljung-Box test proves SARIMA is well-fitted, so the XGBoost win is genuine.'
)

h2("6.10 Inventory Simulation")
body(
    "Using the model's forecasts to simulate a real (s,Q) inventory policy over the "
    "60 test weeks. Four scenarios compared: Naive (constant order), XGBoost forecast, "
    "SARIMA forecast, Perfect Foresight (theoretical optimum)."
)
body("Results:")
bullet("XGBoost saves EUR 273 vs naive policy over 60 weeks (one pharmacy)")
bullet("Fewer stockout events with XGBoost than with SARIMA or naive")
bullet("Perfect Foresight shows the upper bound of possible savings")
analogy(
    "The naive policy is like always buying the same amount of groceries every week "
    "regardless of whether guests are coming. The XGBoost policy is like checking "
    "the guest calendar 6 months out and buying accordingly."
)

h2("6.12 R06 Model — Generalisation")
body(
    "Everything done for R03 (respiratory) was repeated for R06 (antihistamines) "
    "to test whether the framework generalises to related drug categories."
)
bullet("R06 MAPE: 52.07%, R2 = 0.561")
bullet("Framework works — slightly worse performance than R03 because the flu-antihistamine link is weaker than the flu-respiratory link")
bullet("DM test: XGBoost still significantly outperforms SARIMA for R06")
tip('"Is the framework limited to R03?" — No. It generalises to R06 and potentially any seasonal medicine category correlated with respiratory infections. Section 6.12 proves this empirically.')

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7: BUSINESS INTELLIGENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 7: Business Intelligence — So What? The Real-World Value")
body(
    "This chapter translates all the technical work into business language. "
    "It answers: 'OK, the model works, but how does a real company use it?'"
)

h2("7.1 From Reactive to Predictive: The S&OP Integration")
body(
    "Current pharma supply chains run a monthly S&OP (Sales & Operations Planning) cycle. "
    "The model converts this from a backward-looking exercise (what did we sell last month?) "
    "to a forward-looking one (what will we need in 6 months?)."
)
body("The practical annual cycle with this framework:")
bullet("July-August: Australian flu season peaks. Monitor WHO FluNet weekly.")
bullet("August: Model generates 26-week demand forecast for January-February.")
bullet("September: Procurement manager reviews forecast + dashboard alerts.")
bullet("October: Manufacturing batch order placed based on forecast.")
bullet("January-February: Medicines arrive. Stock is pre-positioned for peak demand.")
analogy(
    "It's the difference between a reactive fire department (wait for the fire, then respond) "
    "and a predictive one (use weather forecasts and historical data to pre-position "
    "firetrucks in high-risk areas before fire season starts)."
)

h2("7.2 Operational Integration")
body(
    "The thesis proposes a fully automated production pipeline: "
    "every Sunday, the WHO FluNet API is queried automatically, the new data is appended, "
    "the model retrains (or updates), and the updated 26-week forecast is pushed to the "
    "S&OP planning system. Zero human intervention required for routine updates."
)
body(
    "Total data cost: EUR 0. WHO FluNet is free and public. "
    "The only cost is compute time (the model runs in seconds on any laptop)."
)

h2("7.3 The Streamlit Dashboard")
body(
    "The dashboard is the operational face of the entire pipeline. "
    "It has 13 pages covering every analytical output, from the executive KPI summary "
    "to the SHAP explainability drill-down. Non-technical supply chain managers can "
    "use the Calculadora Predictiva page without understanding XGBoost at all — "
    "they enter current conditions and get a traffic-light risk rating (LOW/MEDIUM/HIGH) "
    "with a specific reorder recommendation."
)
tip(
    '"Can a non-technical person use this?" — Yes, that was an explicit design goal. '
    'The Calculadora Predictiva abstracts away all ML complexity. The manager enters '
    'current stock, lead time, and current flu levels. The system outputs: reorder now '
    'or not, and how much. That\'s actionable information without requiring any '
    'data science knowledge.'
)

h2("7.4 Change Management")
body(
    "Even a perfect model fails if people don't trust or use it. The thesis addresses "
    "this with: (1) SHAP explanations that show WHY the model predicted what it predicted; "
    "(2) confidence intervals that communicate uncertainty honestly; (3) a switching rule "
    "that explicitly acknowledges when the model is unreliable (summer)."
)

h2("7.5 Financial Impact")
body(
    "The EUR 273 saving over 60 weeks from one pharmacy scales dramatically. "
    "A national pharmacy chain with 1,000 locations could save EUR 273,000 from this "
    "alone — just from better timing of orders. Add reduced stockouts (lost sales) "
    "and the figure grows further. The thesis is conservative in its estimates "
    "precisely to avoid overstating the business case."
)

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8: CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("Chapter 8: Conclusions — What Did You Prove?")

h2("Main Findings")
bullet("CONFIRMED: Australian flu data predicts European pharmaceutical demand 26-28 weeks ahead (r=0.70, p<0.001)")
bullet("CONFIRMED: XGBoost with epidemiological features significantly outperforms SARIMA baseline (DM p<0.001)")
bullet("CONFIRMED: Framework generalises to 7 SH countries and 2 drug categories")
bullet("CONFIRMED: Inventory savings are real and measurable (EUR 273 over 60 weeks)")
bullet("CONFIRMED: Seasonal switching rule improves MAPE by 12.85pp (best result: 35.78%)")
bullet("NEGATIVE RESULT: Multi-country ensemble does not improve over Australia alone")

h2("8.1 Study Limitations — Be Honest About These")
body(
    "You must know these cold because the tribunal will probe them. "
    "Good news: you already acknowledge them in the thesis, which is the right approach."
)
bullet("Single pharmacy dataset (Kaggle): results may not generalise to all pharmacy contexts. BUT: AMELI France validates the seasonal pattern at national scale.")
bullet("MAPE of 44%: seems high. BUT: off-season dominates the average. Peak-season performance is much better, and the model's value is in direction and timing, not precision.")
bullet("6-year training window (2014-2019): COVID disrupted patterns post-2019. BUT: the biological mechanism (hemispheric flu lag) is stable across decades.")
bullet("One drug category (R03 primary): BUT Section 6.12 extends to R06, establishing generalisability.")
bullet("No real-time production deployment: this is a research prototype, not a live system. BUT: the dashboard and pipeline are fully functional.")

h2("8.2 Future Work")
bullet("Retrain with post-2020 data once COVID disruption patterns stabilise")
bullet("Extend to other ATC categories (R01 nasal, R05 cough) with same framework")
bullet("Pilot with real pharmacy chain data to test at scale")
bullet("Test alternative ML models (LightGBM, Prophet) as Model B variants")
bullet("Add climate/weather features (cold snaps correlate with flu outbreaks)")

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# LIKELY DEFENSE QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("LIKELY DEFENSE QUESTIONS & ANSWERS")
body(
    "These are the questions most likely to come from your tribunal. "
    "Read these out loud to yourself until the answers feel natural."
)

qa(
    "Why is your MAPE so high at 44%?",
    "MAPE of 44% includes summer weeks where demand is near-zero and any prediction "
    "produces large relative errors mathematically. During peak season (Oct-Apr, when "
    "procurement decisions matter), model error is significantly lower. Furthermore, "
    "with the seasonal switching rule, the overall MAPE drops to 35.78%. The correct "
    "benchmark is not 'perfect accuracy' but 'better than existing methods' — and the "
    "model beats SARIMA by 12.65pp with DM p<0.001."
)
qa(
    "Isn't 302 weeks too little data to train an ML model?",
    "It is a limitation acknowledged in Section 8.1. However, the walk-forward validation "
    "with 48 folds and 192 OOS predictions provides robust evidence. We also incorporated "
    "WHO FluNet data going back to 1997, and the epidemiological lag signal is based on "
    "decades of consistent hemispheric flu seasonality. For a TFG scope, the dataset is "
    "appropriate and honestly validated."
)
qa(
    "Why didn't the ensemble of AU+NZ+Chile improve performance?",
    "This was our expectation too, but the data rejected the hypothesis. Adding more "
    "Southern Hemisphere signals introduces noise because: (1) NZ and Chile have smaller "
    "surveillance networks than Australia, (2) their optimal lags differ from Australia's "
    "(27w and 35w), which creates phase misalignment in the combined signal, and "
    "(3) with 302 training weeks, the model cannot reliably extract signal from "
    "three weakly correlated predictors simultaneously."
)
qa(
    "Is one Kaggle pharmacy representative of all European pharmacies?",
    "We address this specifically in Section 4.1. The AMELI France national dataset "
    "(68 million inhabitants, 11 years) confirms identical seasonal patterns and the "
    "same R03 > R06 relationship. The scale is ~1,000x as expected. This external "
    "validation strongly suggests the seasonal dynamics are consistent across European "
    "markets, even if the absolute demand units are pharmacy-specific."
)
qa(
    "How would this work in a real company?",
    "The pipeline is designed for automation: WHO FluNet data is free and queryable via "
    "public API weekly. The complete pipeline (17 Python scripts) runs end-to-end in "
    "under 10 minutes on a laptop. The Streamlit dashboard runs locally with one command. "
    "A production deployment would require: (1) connecting company sales data as the "
    "target variable, (2) scheduling the pipeline via cron job, (3) integrating "
    "dashboard outputs with the company's S&OP tool."
)
qa(
    "Why XGBoost and not deep learning (LSTM, Transformer)?",
    "Three reasons: (1) With 302 training weeks, deep learning models would overfit — "
    "they need orders of magnitude more data. (2) XGBoost is the industry standard "
    "for structured tabular forecasting and consistently outperforms deep learning on "
    "small tabular datasets. (3) Interpretability: SHAP works natively with XGBoost. "
    "LSTMs are black boxes with no comparable interpretability tool for tabular data."
)
qa(
    "What is the Diebold-Mariano test and why does it matter?",
    "The DM test formally tests whether the difference in forecast accuracy between "
    "two models is statistically significant or could have occurred by chance. "
    "Our DM statistic of 6.23 with p<0.001 means there is less than 0.1% chance "
    "the XGBoost advantage over SARIMA is random. This is the gold standard for "
    "comparing forecast models in the economics and finance literature. "
    "Without it, you cannot claim your model 'wins' — you can only say it happened "
    "to do better on this particular test set."
)
qa(
    "What would you do differently if you started again?",
    "I would source weekly national-level pharmacy sales data from a distributor rather "
    "than relying on a single-pharmacy Kaggle dataset. I would also incorporate climate "
    "data (temperature, humidity) as additional flu-correlated features, and test the "
    "framework on post-COVID 2020-2024 data to verify the lag pattern survives the "
    "disruption period. These are exactly the directions outlined in Section 8.2 as "
    "future work."
)
qa(
    "Why didn't you test more models — LightGBM, Prophet, LSTM, Transformer?",
    "The objective of this thesis was NOT to find the lowest MAPE in a model benchmark "
    "competition. It was to validate a specific scientific hypothesis: that the Australia-"
    "Europe epidemiological lead-lag can be operationalised as a pharmaceutical forecasting "
    "signal. XGBoost was chosen because it is the industry standard for small structured "
    "tabular datasets, and because SHAP provides the interpretability required in regulated "
    "healthcare contexts.\n\n"
    "That said, Prophet (Meta) WAS tested as part of the thesis: MAPE = 48.32%. The "
    "Diebold-Mariano test showed p = 0.228 — meaning XGBoost's 4.16pp advantage over "
    "Prophet is NOT statistically significant at the 5% level. This is an honest finding "
    "documented in Section 6.14. XGBoost is still preferred because of SHAP interpretability, "
    "not because it is categorically better than Prophet on this dataset.\n\n"
    "LightGBM would likely produce similar results to XGBoost on this dataset size — the "
    "marginal gain would not constitute a meaningful scientific contribution. LSTMs and "
    "Transformer-based models require orders of magnitude more data than 302 weekly "
    "observations to avoid overfitting; they are documented as future work in Section 8.2 "
    "for when multi-year national-level data becomes available."
)

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# TECHNICAL GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h1("TECHNICAL GLOSSARY — Plain English Definitions")
body("When the tribunal uses technical terms, you need to understand and respond correctly.")

terms = [
    ("MAPE", "Mean Absolute Percentage Error. Average % by which predictions miss actual values. Lower is better. 35.78% means predictions are off by 35.78% on average."),
    ("MAE", "Mean Absolute Error. Average absolute units by which predictions miss. Easier to interpret than MAPE because it's in the same units as demand."),
    ("R-squared (R²)", "How much of the demand variation the model explains. R²=0.5 means the model explains 50% of variation. R²<0 means the model is worse than just predicting the mean — common in volatile seasonal data."),
    ("XGBoost", "eXtreme Gradient Boosting. A machine learning algorithm that builds hundreds of decision trees, each correcting errors of the previous ones. Industry standard for tabular prediction."),
    ("SARIMA", "Seasonal AutoRegressive Integrated Moving Average. The classical statistical model for seasonal time series. Does NOT use external variables."),
    ("CCF", "Cross-Correlation Function. Measures how correlated two time series are when one is shifted by N weeks (lag). Peak at lag 28w = r=0.70 is the empirical core of this thesis."),
    ("SHAP", "SHapley Additive exPlanations. Tells you how much each feature contributed to a specific prediction. Makes XGBoost interpretable."),
    ("Walk-Forward Validation", "Rolling time-series cross-validation. Train on all past data, predict next N weeks, roll forward. More rigorous than a single train/test split."),
    ("Diebold-Mariano Test", "Statistical test to check if two forecasting models are significantly different in accuracy. DM > 0 means model 1 wins. p < 0.05 = statistically significant win."),
    ("Ljung-Box Test", "Tests if residuals of a time series model are 'white noise' (no remaining pattern). Passing this test means your SARIMA model is well-specified."),
    ("Lead-Lag", "When one time series predicts another with a time delay. AU flu leads EU demand by 28 weeks — Australia is the 'leading indicator'."),
    ("Bullwhip Effect", "Demand signal amplification through a supply chain. Small retail fluctuations → huge factory order swings. Caused by information delay and overreaction."),
    ("S&OP", "Sales & Operations Planning. The monthly business process where sales forecasts are translated into production and inventory plans."),
    ("ATC Code", "Anatomical Therapeutic Chemical code. WHO drug classification system. R03 = respiratory/bronchodilators, R06 = antihistamines."),
    ("(s,Q) Policy", "Inventory reorder policy: when stock drops to level s (reorder point), order quantity Q. Simple and widely used in practice."),
    ("Confidence Interval", "Range within which the actual value is expected to fall X% of the time. An 80% CI means actual demand falls inside the range 80% of weeks."),
    ("Data Leakage", "When future information accidentally influences model training, inflating apparent performance. Avoided here by strict chronological train/test split."),
    ("WHO FluNet", "World Health Organization's global influenza surveillance database. Free, public, updated weekly. The primary data source for the Australian leading indicator."),
    ("Switching Rule", "The hybrid model: use XGBoost for peak season (weeks 1-21, 40-52), use historical seasonal mean for off-season (weeks 22-39). Best MAPE result: 35.78%."),
]

for term, definition in terms:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{term}: ")
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(10.5)
    r2 = p.add_run(definition)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BLACK

pagebreak()


# ══════════════════════════════════════════════════════════════════════════════
# ONE-PAGE SUMMARY (print this and carry it)
# ══════════════════════════════════════════════════════════════════════════════
h1("ONE-PAGE SUMMARY — Print and Carry to the Defense")

h2("What the thesis does")
body("Builds a machine learning system that predicts European respiratory medicine demand 26 weeks in advance using real Australian flu surveillance data from WHO FluNet.")

h2("Why it works")
body("Australian flu season (June-August) precedes European season (December-February) by 28 weeks. CCF r=0.70, p<0.001. This is the epidemiological foundation.")

h2("How it was built")
body("6 research phases: Kaggle only → + Australia FluNet → + Europe FluNet → + AMELI validation → + 7-country validation → + Switching Rule optimisation.")

h2("Key results")
keynum("Best MAPE", "35.78%", "Seasonal Switching Rule (hybrid model)")
keynum("XGBoost vs SARIMA", "DM=6.23, p<0.001", "Statistically significant win")
keynum("Inventory saving", "EUR 273 / 60 weeks", "One pharmacy vs naive policy")
keynum("Countries validated", "7", "All Southern Hemisphere major countries")
keynum("Forward horizon", "26 weeks", "Longer than manufacturing lead time")

h2("Honest limitations")
body("Single pharmacy data. Post-COVID patterns not tested. MAPE includes high-error summer weeks. These are acknowledged and addressed in Section 8.1.")

h2("The one sentence")
p = doc.add_paragraph()
r = p.add_run(
    "Australian flu data, available for free from WHO, predicts European medicine "
    "demand 6 months ahead with r=0.70 correlation — enough to transform reactive "
    "pharmaceutical supply chains into predictive ones."
)
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY


# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
